
import newspaper
import pandas as pd
from docx import Document
import os
from newspaper import Config
import sys
from baiduTrans import BaiduTranslate
import time
from docx.shared import Pt
from docx.shared import RGBColor

#C:\Users\saha\AppData\Local\Temp\.newspaper_scraper\memoized

    
class GetNews(object):
    def __init__(self, bTransBody=False):
        urls = pd.read_excel('config.xls',sheet_name='url')
        self.urlList = urls['url'] 
        keywords = pd.read_excel('config.xls',sheet_name='keyWord')
        self.keywordList = []
        for j in range(0, len(keywords)):
            tmp = keywords.iloc[j]['keywords']     
            arr = tmp.split(';')
            self.keywordList.extend(arr[:-1])
        
        if not os.path.exists('result'):
            os.mkdir('result')
            
        self.news_title = []
        self.news_title_trans = []
        self.news_text = []
        self.news_urls = []
        self.news_keys = []
        self.news_date = []
        self.bTransBody = bTransBody
        self.colletxls = 'result/result_' + time.strftime("%Y%m%d_%H%M%S", time.localtime()) + '.xls'
            
        userInfo = pd.read_excel('config.xls',sheet_name='BaiduTranslate')    
        self.translate = BaiduTranslate('auto','zh',userInfo.iloc[0]['appid'],userInfo.iloc[0]['secretKey'] )
            
        
        
    def readNews(self):        
        user_agent = 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10.15; rv:78.0) Gecko/20100101 Firefox/78.0'
        config = Config()
        config.browser_user_agent = user_agent
        config.fetch_images = False
        config.request_timeout = 20
        config.number_threads = 1
        config.keep_article_html = True
        i = 0
        for url in self.urlList:   
            i = i + 1
            print(str(i) + "  " + url)
            newsSource = newspaper.build(url,config=config)             
            for paper in newsSource.articles:
                try :
                    paper.download()
                    paper.parse()
                    #print('parse')
                    res = self.keyWordMatch(paper.text,paper.title)
                    if len(res)>0:
                        print('hint: ' + paper.title)
                        title_trans = self.translate.BdTrans(paper.title)
                        title_trans = title_trans[0]['dst']
                        self.news_title_trans.append(title_trans)
                        self.news_title.append(paper.title)     # 将新闻题目以列表形式逐一储存
                        self.news_text.append(paper.text)       # 将新闻正文以列表形式逐一储存
                        self.news_urls.append(paper.url)
                        self.news_keys.append(res)
                        pDate=''
                        if paper.publish_date is not None:
                            pDate = paper.publish_date.strftime('%Y-%m-%d')
                            print(pDate)
                        self.news_date.append(pDate)     
                        body_trans = []
                        if self.bTransBody:
                            body_trans = self.translate.BdTrans(paper.text)
                        self.saveFile(pDate,paper.url,res,paper.title,paper.text,title_trans,body_trans)
                        self.sumExcel()
                except:
                    print("Unexpected error:", sys.exc_info()[0])
                    continue                

        print('finished!')
        
    def sumExcel(self):
        data = pd.DataFrame({'title':self.news_title,'title_trans':self.news_title_trans,'key':self.news_keys,'date':self.news_date,'url':self.news_urls})
        data.to_excel(self.colletxls,sheet_name='result')
        
                
    def keyWordMatch(self, text,title):
        for keyword in self.keywordList:
            if keyword in title:
                return keyword
            if keyword in text:
                return keyword
        
        return ''
    
    def saveFile(self,pb_date,url,keyWord,title,body,title_trans,body_trans):
        document = Document()
        document.add_heading(title + '\n' + title_trans, 0)
        p=document.add_paragraph('')
        p.add_run('url: ' + url + '\n').bold = True
        p.add_run('publish_date: ' + pb_date + '\n').bold = True
        p.add_run('match key_word: ' + keyWord + '\n').bold = True
        if len(body_trans) == 0:
            p.add_run(body)
        else:
            for detail in body_trans:
                p.add_run(detail['src'])
                p.add_run('\n')
                p_dst=p.add_run(detail['dst'])
                p_dst.font.size = Pt(8)
                p_dst.font.italic = True
                p_dst.font.color.rgb = RGBColor(0xCC, 0x41, 0x25)
                p.add_run('\n')
        #document.add_page_break() 增加分页符
        document.save('result/'+ title +'.docx')
        
    def docTest(self):
        document = Document()
        document.add_heading('title', 0)
        p=document.add_paragraph('')
        p.add_run('publish_date: \n').bold = True
        p.add_run('match key_word: \n').italic = True
        
        p1=p.add_run('测试一下')
        p1.font.size = Pt(8)
        p1.font.italic = True
        p1.font.color.rgb = RGBColor(0xCC, 0x41, 0x25)

        document.save('result/test.docx')
        
        
            
if __name__ == '__main__':
    ins = GetNews(True)
    ins.readNews()
    #ins.docTest()
